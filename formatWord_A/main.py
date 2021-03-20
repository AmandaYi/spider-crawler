import os
from docx import Document
from win32com import client as wc #导入模块
import shutil
from docx.shared import Pt,RGBColor
from docx.oxml.ns import qn
from win32com import client as wc  

targetPath = os.getcwd() + "./targetPath/"
docxPath = os.getcwd() + "./docxPath/" # 即将处理的全部文件
resultPath = os.getcwd() + "./resultPath/"

print("1、拷贝工作*****为防止损坏原始文件，开始拷贝文件到待处理目录！")
files = []
for file in os.listdir(targetPath):
    if file.endswith(".doc"):
       files.append(file)
    if file.endswith(".docx"):
        # 移动文件到docxPath
        print("拷贝工作 -> 拷贝正常文件docx [" + file + "] 到待处理目录")
        path = shutil.copy(targetPath + file, docxPath + file)

word = wc.Dispatch("Word.Application") 

print("2、修正工作开始***********************************************************************开始修正,规则doc -> docx")
for file in files:
    doc = word.Documents.Open(targetPath + file) #打开word文件
    doc.SaveAs(docxPath + file + "x", 12) #另存为后缀为".docx"的文件，其中参数12指docx文件
    doc.Close() #关闭原来word文件
    print("修正工作 -> 文件 [", file + "] 是doc结尾，---> 默认修改为docx结尾进行处理")
word.Quit()

# 开始循环解析每一个项目URL，得到的内容赋给一个新值
# 开始处理数据
print("3、正式进入处理阶段********************************************，请不要中途退出或者强退！")
print("请回车，开始执行处理")
s = input("")

fileList = os.listdir(docxPath)
fileListLen = len(fileList)
print("正式处理工作 -> 文件数量：" + str(fileListLen))
for fileIndex in range(len(fileList)):
    file = fileList[fileIndex]
    # 防止报错，如果是docx的才处理，否则跳过此文件
    if file.endswith(".doc"): 
        continue
    print("正式处理工作 -> STR------开始处理文件" + str(fileIndex + 1) + ": [" + file + "] STR-----")
    document = Document(docxPath + file)
    # 字号‘八号’对应磅值5
    # 字号‘七号’对应磅值5.5
    # 字号‘小六’对应磅值6.5
    # 字号‘六号’对应磅值7.5
    # 字号‘小五’对应磅值9
    # 字号‘五号’对应磅值10.5
    # 字号‘小四’对应磅值12
    # 字号‘四号’对应磅值14
    # 字号‘小三’对应磅值15
    # 字号‘三号’对应磅值16
    # 字号‘小二’对应磅值18
    # 字号‘二号’对应磅值22
    # 字号‘小一’对应磅值24
    # 字号‘一号’对应磅值26
    # 字号‘小初’对应磅值36
    # 字号‘初号’对应磅值42
    
    print("正式处理工作 -> 文件 [" + file + "]")
    print("去掉页眉页脚")
    for section in document.sections:
        for sectionHeaderParagraphs in section.header.paragraphs:
            sectionHeaderParagraphs.text = ""
        for sectionFooterParagraphs in section.footer.paragraphs:
            sectionFooterParagraphs.text = ""
    
    print("开始设置全文字体为宋体")
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    print("正式处理工作 -> 文件 [" + file + "] , 开始设置---规则标题大小一号， 全文大小四号，请等待！！！")
    for index in range(len(document.paragraphs)):
        paragraph = document.paragraphs[index]
        for runIndex in range(len(paragraph.runs)):
            if index == 0 & runIndex == 0:
                # 第一自然段是一号字体
                paragraph.runs[runIndex].font.size = Pt(26)
            else:
                paragraph.runs[runIndex].font.size = Pt(14)
            # 全局字体是宋体
            paragraph.runs[runIndex].font.name = u'宋体'
    document.save(resultPath + file)
    print("正式处理工作 -> END------处理完成文件" + str(fileIndex + 1) + ": [" + file + "] END-----")

print("------------处理结果请查看文件夹" + resultPath + "------------")


    